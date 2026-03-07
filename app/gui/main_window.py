"""
Digital Alchemist X-1 - Main GUI Window
Dark theme, 5-model tabs, real-time word count, background processing.
Copy/Paste/Cut/Select-All fully enabled via right-click menu and Ctrl shortcuts.
Left panel is fully scrollable.
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import threading
import os
import sys
from datetime import datetime
from typing import Dict, Optional

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(
    os.path.abspath(__file__)))))

from app.core.text_processor import TextProcessor, STYLE_CONFIGS

# ── Colors ─────────────────────────────────────────────────────────
BG_DARK    = '#1a1a2e'
BG_PANEL   = '#16213e'
BG_INPUT   = '#0f3460'
BG_DARKER  = '#0d0d1a'
ACCENT     = '#e94560'
ACCENT2    = '#c73652'
TEXT_WHITE = '#eaeaea'
TEXT_GRAY  = '#a0a0b0'
SUCCESS    = '#4caf50'
TAB_ACTIVE = '#e94560'
TAB_IDLE   = '#16213e'

STYLE_LABELS = {
    1: ('[A]', 'Model 1: Analytical-'),
    2: ('[E]', 'Model 2: Empirical-Sc'),
    3: ('[C]', 'Model 3: Critical-Argum'),
    4: ('[T]', 'Model 4: Theoretical-Co'),
    5: ('[D]', 'Model 5: Descriptive-Ex'),
}


def _attach_context_menu(widget):
    """Attach right-click context menu with Copy/Cut/Paste/Select All."""
    menu = tk.Menu(widget, tearoff=0,
                   bg='#16213e', fg='#eaeaea',
                   activebackground='#e94560',
                   activeforeground='#ffffff',
                   font=('Segoe UI', 9))

    def _copy():
        try:
            widget.event_generate('<<Copy>>')
        except Exception:
            pass

    def _cut():
        try:
            widget.event_generate('<<Cut>>')
        except Exception:
            pass

    def _paste():
        try:
            widget.event_generate('<<Paste>>')
        except Exception:
            pass

    def _select_all():
        try:
            widget.tag_add('sel', '1.0', 'end')
            widget.mark_set('insert', '1.0')
            widget.see('insert')
        except Exception:
            try:
                widget.select_range(0, 'end')
            except Exception:
                pass

    menu.add_command(label='Copy        Ctrl+C',  command=_copy)
    menu.add_command(label='Cut         Ctrl+X',  command=_cut)
    menu.add_command(label='Paste       Ctrl+V',  command=_paste)
    menu.add_separator()
    menu.add_command(label='Select All  Ctrl+A',  command=_select_all)

    def _show(event):
        try:
            menu.tk_popup(event.x_root, event.y_root)
        finally:
            menu.grab_release()

    widget.bind('<Button-3>', _show)

    # Ensure standard Ctrl shortcuts work
    widget.bind('<Control-a>', lambda e: _select_all())
    widget.bind('<Control-A>', lambda e: _select_all())
    widget.bind('<Control-c>', lambda e: _copy())
    widget.bind('<Control-C>', lambda e: _copy())
    widget.bind('<Control-x>', lambda e: _cut())
    widget.bind('<Control-X>', lambda e: _cut())
    widget.bind('<Control-v>', lambda e: _paste())
    widget.bind('<Control-V>', lambda e: _paste())


class MainWindow(tk.Tk):

    def __init__(self):
        super().__init__()
        self.title(
            'Digital Alchemist X-1  --  Academic AI-Humanization Engine')
        self.geometry('1280x820')
        self.minsize(1000, 650)
        self.configure(bg=BG_DARK)

        self._results: Dict[int, str] = {}
        self._active_tab: int = 1
        self._processing: bool = False
        self._processor: Optional[TextProcessor] = None
        self._thread: Optional[threading.Thread] = None
        self._style_checks: Dict[int, tk.BooleanVar] = {}

        self._build_ui()
        self._update_word_count()

    # ──────────────────────────────────────────────────────────────
    #  BUILD UI
    # ──────────────────────────────────────────────────────────────

    def _build_ui(self):
        # Header bar
        hdr = tk.Frame(self, bg=BG_DARKER, height=44)
        hdr.pack(fill='x', side='top')
        hdr.pack_propagate(False)
        tk.Label(
            hdr,
            text='Digital Alchemist X-1  --  Academic AI-Humanization Engine',
            bg=BG_DARKER, fg=ACCENT,
            font=('Segoe UI', 13, 'bold')
        ).pack(side='left', padx=16, pady=10)
        tk.Label(
            hdr,
            text='Transform AI text  \u2192  Undetectable Scholarly Prose  |  5 Academic Models',
            bg=BG_DARKER, fg=TEXT_GRAY,
            font=('Segoe UI', 9)
        ).pack(side='right', padx=16)

        # Main layout
        main = tk.Frame(self, bg=BG_DARK)
        main.pack(fill='both', expand=True, padx=6, pady=6)

        # ── Left panel with scrollbar ──────────────────────────────
        left_outer = tk.Frame(main, bg=BG_DARK, width=362)
        left_outer.pack(side='left', fill='y')
        left_outer.pack_propagate(False)

        left_sb = tk.Scrollbar(left_outer, orient='vertical',
                                bg=BG_DARKER, troughcolor=BG_DARKER,
                                activebackground=ACCENT,
                                highlightthickness=0, bd=0)
        left_sb.pack(side='right', fill='y')

        self._left_canvas = tk.Canvas(
            left_outer, bg=BG_DARK, bd=0,
            highlightthickness=0,
            yscrollcommand=left_sb.set)
        self._left_canvas.pack(side='left', fill='both', expand=True)

        left_sb.config(command=self._left_canvas.yview)

        # Inner frame that holds all left-panel widgets
        self._left_inner = tk.Frame(self._left_canvas, bg=BG_DARK)
        self._left_canvas_win = self._left_canvas.create_window(
            (0, 0), window=self._left_inner, anchor='nw')

        # Resize canvas window width when canvas resizes
        def _on_canvas_resize(event):
            self._left_canvas.itemconfig(
                self._left_canvas_win, width=event.width)
        self._left_canvas.bind('<Configure>', _on_canvas_resize)

        # Update scroll region when inner frame changes
        def _on_inner_resize(event):
            self._left_canvas.configure(
                scrollregion=self._left_canvas.bbox('all'))
        self._left_inner.bind('<Configure>', _on_inner_resize)

        # Mouse wheel scroll on left panel
        def _on_mousewheel(event):
            self._left_canvas.yview_scroll(
                int(-1 * (event.delta / 120)), 'units')
        self._left_canvas.bind_all('<MouseWheel>', _on_mousewheel)

        # Right panel
        right = tk.Frame(main, bg=BG_DARK)
        right.pack(side='left', fill='both', expand=True, padx=(6, 0))

        self._build_left(self._left_inner)
        self._build_right(right)

    def _lf(self, parent, title):
        return tk.LabelFrame(
            parent, text=title,
            bg=BG_PANEL, fg=ACCENT,
            font=('Segoe UI', 9, 'bold'),
            bd=1, relief='flat',
            padx=4, pady=4)

    def _build_left(self, parent):
        # ── Input Source ──────────────────────────────────────────
        src = self._lf(parent, 'INPUT SOURCE')
        src.pack(fill='x', padx=4, pady=(4, 2))

        self._file_var = tk.StringVar()
        fr = tk.Frame(src, bg=BG_PANEL)
        fr.pack(fill='x', pady=2)
        self._file_entry = tk.Entry(
            fr, textvariable=self._file_var,
            bg=BG_INPUT, fg=TEXT_WHITE,
            insertbackground=TEXT_WHITE,
            relief='flat', font=('Segoe UI', 9))
        self._file_entry.pack(side='left', fill='x', expand=True, ipady=4)
        _attach_context_menu(self._file_entry)
        tk.Button(fr, text='X', bg=BG_INPUT, fg=TEXT_GRAY,
                  relief='flat', font=('Segoe UI', 8), width=2,
                  command=self._clear_file).pack(side='left', padx=1)
        tk.Button(fr, text='Browse', bg=ACCENT, fg=TEXT_WHITE,
                  relief='flat', font=('Segoe UI', 9, 'bold'),
                  cursor='hand2',
                  command=self._browse_file).pack(side='left', padx=2)
        tk.Label(src,
                 text='Accepts: .docx  |  .pdf  |  .txt  (up to 25,000 words)',
                 bg=BG_PANEL, fg=TEXT_GRAY,
                 font=('Segoe UI', 8)).pack(anchor='w')
        tk.Label(src,
                 text='-- OR paste text directly below --',
                 bg=BG_PANEL, fg=TEXT_GRAY,
                 font=('Segoe UI', 8, 'italic')).pack(anchor='w')

        # ── Text Input ────────────────────────────────────────────
        txt = self._lf(parent, 'TEXT INPUT  (paste here)')
        txt.pack(fill='x', padx=4, pady=2)

        self._input_text = scrolledtext.ScrolledText(
            txt, bg=BG_INPUT, fg=TEXT_WHITE,
            insertbackground=TEXT_WHITE,
            font=('Segoe UI', 10), wrap='word',
            relief='flat', bd=0,
            height=12,
            undo=True)
        self._input_text.pack(fill='x', expand=False)
        self._input_text.bind('<KeyRelease>',
                               lambda e: self._update_word_count())
        _attach_context_menu(self._input_text)

        self._wc_label = tk.Label(
            txt, text='Words: 0',
            bg=BG_PANEL, fg=TEXT_GRAY,
            font=('Segoe UI', 8))
        self._wc_label.pack(anchor='e', pady=(2, 0))

        # ── Model Selection ───────────────────────────────────────
        sel = self._lf(parent, 'SELECT OUTPUT MODELS')
        sel.pack(fill='x', padx=4, pady=2)

        for sid, cfg in STYLE_CONFIGS.items():
            var = tk.BooleanVar(value=True)
            self._style_checks[sid] = var
            code, short = STYLE_LABELS[sid]
            row = tk.Frame(sel, bg=BG_PANEL)
            row.pack(fill='x', pady=1)
            tk.Checkbutton(
                row, variable=var,
                bg=BG_PANEL, fg=TEXT_WHITE,
                selectcolor=BG_INPUT,
                activebackground=BG_PANEL,
                activeforeground=TEXT_WHITE
            ).pack(side='left')
            tk.Label(row,
                     text=f'{code} {cfg["name"]}',
                     bg=BG_PANEL, fg=TEXT_WHITE,
                     font=('Segoe UI', 9, 'bold')
                     ).pack(side='left')
            tk.Label(row,
                     text=f'  {cfg["description"][:38]}',
                     bg=BG_PANEL, fg=TEXT_GRAY,
                     font=('Segoe UI', 8)
                     ).pack(side='left')

        btn_row = tk.Frame(sel, bg=BG_PANEL)
        btn_row.pack(fill='x', pady=(4, 0))
        tk.Button(btn_row, text='Select All',
                  bg=BG_INPUT, fg=TEXT_WHITE,
                  relief='flat', font=('Segoe UI', 9),
                  command=lambda: [v.set(True)
                                   for v in self._style_checks.values()]
                  ).pack(side='left', padx=4)
        tk.Button(btn_row, text='Select None',
                  bg=BG_INPUT, fg=TEXT_WHITE,
                  relief='flat', font=('Segoe UI', 9),
                  command=lambda: [v.set(False)
                                   for v in self._style_checks.values()]
                  ).pack(side='left')

        # ── Process ───────────────────────────────────────────────
        proc = self._lf(parent, 'PROCESS')
        proc.pack(fill='x', padx=4, pady=2)

        self._progress = ttk.Progressbar(proc, mode='determinate',
                                          maximum=100)
        self._progress.pack(fill='x', pady=(2, 0))

        self._prog_label = tk.Label(proc, text='',
                                     bg=BG_PANEL, fg=TEXT_GRAY,
                                     font=('Segoe UI', 8))
        self._prog_label.pack(anchor='w')

        self._start_btn = tk.Button(
            proc, text='START PROCESSING',
            bg=ACCENT, fg=TEXT_WHITE,
            font=('Segoe UI', 10, 'bold'),
            relief='flat', cursor='hand2', pady=8,
            command=self._start_processing)
        self._start_btn.pack(fill='x', pady=(4, 2))

        self._stop_btn = tk.Button(
            proc, text='STOP',
            bg=ACCENT2, fg=TEXT_WHITE,
            font=('Segoe UI', 10, 'bold'),
            relief='flat', cursor='hand2', pady=6,
            command=self._stop_processing,
            state='disabled')
        self._stop_btn.pack(fill='x', pady=(0, 4))

        # ── Log ───────────────────────────────────────────────────
        log = self._lf(parent, 'PROCESSING LOG')
        log.pack(fill='x', padx=4, pady=(2, 8))

        self._log = scrolledtext.ScrolledText(
            log, bg=BG_DARKER, fg=SUCCESS,
            font=('Consolas', 8), wrap='word',
            height=6, relief='flat', state='disabled')
        self._log.pack(fill='x', expand=False)
        _attach_context_menu(self._log)

    def _build_right(self, parent):
        # ── Tab bar ───────────────────────────────────────────────
        tab_bar = tk.Frame(parent, bg=BG_DARKER)
        tab_bar.pack(fill='x', side='top')

        self._tab_btns: Dict[int, tk.Button] = {}
        for sid in range(1, 6):
            code, short = STYLE_LABELS[sid]
            btn = tk.Button(
                tab_bar,
                text=f'{code} {short}',
                bg=TAB_IDLE, fg=TEXT_GRAY,
                font=('Segoe UI', 9),
                relief='flat', bd=0,
                padx=10, pady=7,
                cursor='hand2',
                command=lambda s=sid: self._switch_tab(s))
            btn.pack(side='left', padx=1, pady=1)
            self._tab_btns[sid] = btn

        tk.Button(
            tab_bar, text='Save \u25be',
            bg=BG_INPUT, fg=TEXT_WHITE,
            font=('Segoe UI', 9),
            relief='flat', padx=10, pady=7,
            cursor='hand2',
            command=self._save_output
        ).pack(side='right', padx=4, pady=1)

        # ── Output header ──────────────────────────────────────────
        hdr_frame = tk.Frame(parent, bg=BG_PANEL)
        hdr_frame.pack(fill='x', side='top')

        self._out_header = tk.Label(
            hdr_frame, text='',
            bg=BG_PANEL, fg=TEXT_WHITE,
            font=('Segoe UI', 9), anchor='w', padx=8)
        self._out_header.pack(side='left', fill='x', expand=True, pady=4)

        self._out_stats = tk.Label(
            hdr_frame, text='0 words  |  0 chars',
            bg=BG_PANEL, fg=TEXT_GRAY,
            font=('Segoe UI', 9), anchor='e', padx=8)
        self._out_stats.pack(side='right', pady=4)

        # ── Output text ────────────────────────────────────────────
        self._output_text = scrolledtext.ScrolledText(
            parent,
            bg=BG_INPUT, fg=TEXT_WHITE,
            font=('Segoe UI', 11), wrap='word',
            relief='flat', bd=0,
            state='normal')
        self._output_text.pack(fill='both', expand=True, padx=4, pady=4)
        self._output_text.configure(state='disabled')
        _attach_context_menu(self._output_text)

        self._switch_tab(1)

    # ──────────────────────────────────────────────────────────────
    #  TAB SWITCHING
    # ──────────────────────────────────────────────────────────────

    def _switch_tab(self, sid: int):
        self._active_tab = sid
        for s, btn in self._tab_btns.items():
            if s == sid:
                btn.configure(bg=TAB_ACTIVE, fg=TEXT_WHITE,
                               font=('Segoe UI', 9, 'bold'))
            else:
                btn.configure(bg=TAB_IDLE, fg=TEXT_GRAY,
                               font=('Segoe UI', 9))

        cfg  = STYLE_CONFIGS.get(sid, {})
        name = cfg.get('name', '')
        desc = cfg.get('description', '')
        self._out_header.configure(
            text=f'Model {sid}  |  {name}  --  {desc}')

        content = self._results.get(sid, '')
        self._output_text.configure(state='normal')
        self._output_text.delete('1.0', 'end')
        if content:
            self._output_text.insert('1.0', content)
        self._output_text.configure(state='disabled')

        wc = len(content.split()) if content else 0
        cc = len(content) if content else 0
        self._out_stats.configure(text=f'{wc} words  |  {cc} chars')

    # ──────────────────────────────────────────────────────────────
    #  WORD COUNT
    # ──────────────────────────────────────────────────────────────

    def _update_word_count(self):
        txt = self._input_text.get('1.0', 'end').strip()
        wc  = len(txt.split()) if txt else 0
        self._wc_label.configure(text=f'Words: {wc}')

    # ──────────────────────────────────────────────────────────────
    #  FILE OPERATIONS
    # ──────────────────────────────────────────────────────────────

    def _browse_file(self):
        path = filedialog.askopenfilename(
            filetypes=[
                ('Supported files', '*.docx *.txt *.pdf'),
                ('Word documents',  '*.docx'),
                ('Text files',      '*.txt'),
                ('PDF files',       '*.pdf'),
                ('All files',       '*.*'),
            ])
        if not path:
            return
        self._file_var.set(path)
        self._load_file(path)

    def _load_file(self, path: str):
        ext = os.path.splitext(path)[1].lower()
        try:
            if ext == '.txt':
                with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
            elif ext == '.docx':
                try:
                    import docx
                    doc     = docx.Document(path)
                    content = '\n\n'.join(
                        p.text for p in doc.paragraphs if p.text.strip())
                except ImportError:
                    messagebox.showerror(
                        'Missing Package',
                        'python-docx not installed.\n'
                        'Run: pip install python-docx')
                    return
            elif ext == '.pdf':
                messagebox.showinfo(
                    'PDF Not Supported',
                    'PDF reading is not supported.\n'
                    'Please copy-paste the text manually.')
                return
            else:
                messagebox.showerror(
                    'Unsupported',
                    f'File type "{ext}" is not supported.')
                return

            wc = len(content.split())
            if wc > 25000:
                messagebox.showwarning(
                    'Too Long',
                    f'File has {wc} words (max 25,000).\n'
                    'Only first 25,000 words will be loaded.')
                content = ' '.join(content.split()[:25000])

            self._input_text.configure(state='normal')
            self._input_text.delete('1.0', 'end')
            self._input_text.insert('1.0', content)
            self._update_word_count()
            self._log_msg(
                f'[+] Loaded: {os.path.basename(path)} ({wc} words)')

        except Exception as e:
            messagebox.showerror('Error', f'Could not load file:\n{e}')

    def _clear_file(self):
        self._file_var.set('')

    def _save_output(self):
        content = self._results.get(self._active_tab, '')
        if not content:
            messagebox.showinfo('No Output',
                                'No output to save for this model.')
            return
        style_name = STYLE_CONFIGS[self._active_tab]['name']
        path = filedialog.asksaveasfilename(
            defaultextension='.txt',
            filetypes=[('Text file', '*.txt'),
                       ('Word document', '*.docx')],
            initialfile=(f'DigitalAlchemist_Output_Style'
                         f'{self._active_tab}_{style_name}.txt'))
        if not path:
            return
        try:
            if path.endswith('.docx'):
                try:
                    import docx
                    doc = docx.Document()
                    for para in content.split('\n\n'):
                        if para.strip():
                            doc.add_paragraph(para.strip())
                    doc.save(path)
                except ImportError:
                    messagebox.showerror(
                        'Missing Package',
                        'python-docx not installed.\n'
                        'Run: pip install python-docx')
                    return
            else:
                with open(path, 'w', encoding='utf-8') as f:
                    f.write(content)
            self._log_msg(f'[+] Saved: {os.path.basename(path)}')
            messagebox.showinfo('Saved', f'File saved:\n{path}')
        except Exception as e:
            messagebox.showerror('Save Error', f'Could not save:\n{e}')

    # ──────────────────────────────────────────────────────────────
    #  PROCESSING
    # ──────────────────────────────────────────────────────────────

    def _start_processing(self):
        text = self._input_text.get('1.0', 'end').strip()
        if not text:
            messagebox.showwarning('No Input',
                                   'Please enter or paste text first.')
            return
        selected = [sid for sid, var in self._style_checks.items()
                    if var.get()]
        if not selected:
            messagebox.showwarning('No Models',
                                   'Please select at least one model.')
            return

        self._processing = True
        self._results    = {}
        self._start_btn.configure(state='disabled')
        self._stop_btn.configure(state='normal')
        self._progress['value'] = 0
        self._prog_label.configure(text='')
        self._log_msg('[+] Processing started...')

        self._thread = threading.Thread(
            target=self._run_processing,
            args=(text, selected),
            daemon=True)
        self._thread.start()

    def _run_processing(self, text: str, selected: list):
        def cb(msg: str, pct: int):
            self.after(0, lambda m=msg, p=pct: self._on_progress(m, p))

        self._processor = TextProcessor(progress_callback=cb)

        try:
            from langdetect import detect
            try:
                src_lang = detect(text[:600])
            except Exception:
                src_lang = 'en'
        except ImportError:
            src_lang = 'en'

        n_sel = len(selected)

        for i, sid in enumerate(selected):
            if not self._processing:
                break
            self.after(0, lambda s=sid: self._log_msg(
                f'[+] -> Style {s}: {STYLE_CONFIGS[s]["name"]}'))
            try:
                result = self._processor.process_single_style(
                    text, sid, src_lang)
                self._results[sid] = result
                self.after(0, lambda s=sid: self._switch_tab(s))
                self.after(0, lambda s=sid, r=result: self._log_msg(
                    f'[+] Style {s} done: {len(r.split())} words'))
            except Exception as e:
                self._results[sid] = f'[Processing error: {e}]'
                self.after(0, lambda s=sid, err=str(e): self._log_msg(
                    f'[!] Style {s} error: {err}'))

            pct = int((i + 1) / n_sel * 100)
            self.after(0, lambda p=pct: self._progress.configure(value=p))

        self.after(0, self._on_complete)

    def _stop_processing(self):
        self._processing = False
        if self._processor:
            self._processor.stop()
        self._log_msg('[!] Stopping...')

    def _on_progress(self, msg: str, pct: int):
        self._prog_label.configure(text=msg)
        if pct > 0:
            self._progress['value'] = pct

    def _on_complete(self):
        self._processing = False
        self._start_btn.configure(state='normal')
        self._stop_btn.configure(state='disabled')
        self._progress['value'] = 100
        self._prog_label.configure(text='Processing complete!')
        self._log_msg('[+] All done.')
        if self._results:
            best = (self._active_tab if self._active_tab in self._results
                    else min(self._results.keys()))
            self._switch_tab(best)

    # ──────────────────────────────────────────────────────────────
    #  LOG
    # ──────────────────────────────────────────────────────────────

    def _log_msg(self, msg: str):
        ts   = datetime.now().strftime('%H:%M:%S')
        line = f'[{ts}]  {msg}\n'
        self._log.configure(state='normal')
        self._log.insert('end', line)
        self._log.see('end')
        self._log.configure(state='disabled')
